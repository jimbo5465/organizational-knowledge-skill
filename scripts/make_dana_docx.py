# -*- coding: utf-8 -*-
import sys
import docx
from docx import Document
from docx.shared import Pt, Mm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

out = sys.argv[1] if len(sys.argv) > 1 else r"C:\Users\shaterian_m\Documents\Default Project\knowledge creation\پیش‌نویس DANA - پیشنهاد تغییر پوشش مخازن.docx"
doc = Document()

# Page: A4, margins 18mm
sec = doc.sections[0]
sec.page_width, sec.page_height = Mm(210), Mm(297)
sec.right_margin = sec.left_margin = Mm(18)
sec.top_margin = sec.bottom_margin = Mm(18)

DARK = RGBColor(0x0B, 0x3D, 0x62)
BLACK = RGBColor(0x1A, 0x1A, 0x1A)
GREY = RGBColor(0x33, 0x33, 0x33)
HEADC = RGBColor(0x0b, 0x3d, 0x62)

FONT_HEAD = "B Titr"
FONT_BODY = "B Nazanin"


def set_font(run, font, size=11, bold=False, color=BLACK):
    run.font.name = font
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    for a in ("ascii", "hAnsi", "cs", "eastAsia"):
        rFonts.set(qn("w:" + a), font)
    rtl = OxmlElement("w:rtl")
    rtl.set(qn("w:val"), "1")
    rPr.append(rtl)


PSEP = ("w:pStyle", "w:keepNext", "w:keepLines", "w:pageBreakBefore",
        "w:framePr", "w:widowControl", "w:numPr", "w:suppressLineNumbers",
        "w:pBdr", "w:shd", "w:tabs", "w:suppressAutoHyphens", "w:kinsoku",
        "w:wordWrap", "w:overflowPunct", "w:topLinePunct", "w:autoSpaceDE",
        "w:autoSpaceDN", "w:bidi", "w:adjustRightInd", "w:snapToGrid",
        "w:spacing", "w:ind", "w:contextualSpacing", "w:mirrorIndents",
        "w:suppressOverlap", "w:jc", "w:textDirection", "w:textAlignment",
        "w:textboxTightWrap", "w:outlineLvl", "w:divId", "w:cnfStyle",
        "w:rPr", "w:sectPr", "w:pPrChange")


def _insert(pPr, elm, tag):
    if pPr.find(qn(tag)) is not None:
        return
    idx = PSEP.index(tag)
    for later in PSEP[idx + 1:]:
        found = pPr.find(qn(later))
        if found is not None:
            found.addprevious(elm)
            return
    pPr.append(elm)


def set_rtl(p):
    pPr = p._p.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    bidi.set(qn("w:val"), "1")
    _insert(pPr, bidi, "w:bidi")


def set_bottom_border(p, color="0b3d62", size=6):
    pPr = p._p.get_or_add_pPr()
    pbdr = pPr.find(qn("w:pBdr"))
    if pbdr is None:
        pbdr = OxmlElement("w:pBdr")
        _insert(pPr, pbdr, "w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), color)
    pbdr.append(bottom)


def para(text, font=FONT_BODY, size=11, bold=False, align=WD_ALIGN_PARAGRAPH.RIGHT,
         color=BLACK, space_after=3, space_before=0):
    p = doc.add_paragraph()
    set_rtl(p)
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    run = p.add_run(text)
    set_font(run, font, size=size, bold=bold, color=color)
    return p


def heading(text, size=12.5, border=False):
    p = para(text, font=FONT_HEAD, size=size, bold=True, color=DARK,
             space_after=4, space_before=10)
    if border:
        set_bottom_border(p)
    return p


def set_table_rtl(t):
    tblPr = t._tbl.tblPr
    bidi = OxmlElement("w:bidiVisual")
    bidi.set(qn("w:val"), "1")
    tblPr.append(bidi)
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tblPr.append(layout)


def make_table(rows, widths, headfont=FONT_HEAD):
    t = doc.add_table(rows=len(rows), cols=len(rows[0]))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.RIGHT
    t.autofit = False
    for i, r in enumerate(rows):
        cells = t.rows[i].cells
        for j, txt in enumerate(r):
            cell = cells[j]
            cell.width = widths[j]
            p = cell.paragraphs[0]
            set_rtl(p)
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            is_head = (i == 0)
            run = p.add_run(txt)
            set_font(run, headfont if is_head else FONT_BODY,
                     size=10, bold=is_head,
                     color=HEADC if is_head else BLACK)
    set_table_rtl(t)
    return t


# Title
para("پیش‌نویس ثبت دانش در DANA", font=FONT_HEAD, size=16, bold=True,
     align=WD_ALIGN_PARAGRAPH.CENTER, color=DARK, space_after=8)
set_bottom_border(doc.paragraphs[-1], color="0b3d62", size=12)

# Info
heading("اطلاعات ثبت", border=True)
para("نوع دانش: پیشنهاد")
para("وضعیت QA: نیازمند بازبینی")
para("بازبینی اپراتور: الزامی")

# Content
heading("محتوا", border=True)
para("عنوان پیشنهاد", bold=True)
para("تغییر سیستم پوشش مخازن دفنی دیلوج و استوریج کولینگ کمکی از قیری داغ (Bituminous) به اپوکسی کولتار (Coal Tar Epoxy)")
para("وضعیت فعلی", bold=True)
para("طبق پروسیجر پروژه (MD1-TU-00-GN-W-00-PW2-003)، پوشش اصلی مخازن دفنی دیلوج و استوریج کولینگ کمکی (ACS) نیروگاه فردوسی از نوع قیری داغ است. این سیستم در شرایط فعلی پروژه با محدودیت‌های اجرایی، کمبود نیروی ماهر و تجهیزات مناسب مواجه است.")
para("بهبود پیشنهادی", bold=True)
para("جایگزینی پوشش قیری داغ با اپوکسی کولتار بدون حلال (Solvent-Free Coal Tar Epoxy)؛ نمونه پیشنهادی RTB-1319 شرکت روناس با مواد جامد ۹۵–۹۹٪، مقاومت بالا در برابر رطوبت/خاک/محیط قلیایی، اجرای ساده‌تر (Airless Spray، غلتک، قلم)، کنترل ضخامت دقیق‌تر و بدون نیاز به گرم‌کردن.")
para("اثر پیاده‌سازی (مورد انتظار)", bold=True)
para("افزایش دوام و عمر سرویس (۲۰–۳۰ سال در برابر ۵–۱۰ سال)، چسبندگی قوی‌تر، اجرای سریع‌تر و ایمن‌تر، کاهش هزینه نگهداری در طول عمر پروژه، کیفیت یکنواخت و کاهش آسیب در Backfilling. (اثر مورد انتظار — هنوز تأیید نشده)")
para("نتایج پیاده‌سازی", bold=True)
para("ارائه نشده (پیاده‌سازی نشده است)")

# Metadata table
heading("فراداده", border=True)
meta = [
    ("فیلد", "مقدار"),
    ("شجره دانش", "[پیشنهادی] Design and Engineering > Process Engineering > Painting and Coating — ورودی اپراتور الزامی"),
    ("کمیته تخصصی", "کمیته پیشنهادات مدیریت پروژه — ورودی اپراتور الزامی"),
    ("بذر پیشنهاد", "[اختیاری - ارائه نشده]"),
    ("اثر پیاده‌سازی", "مطابق بخش محتوا (اثر مورد انتظار)"),
    ("همکاران", "[اختیاری - ارائه نشده]"),
    ("سطح دسترسی", "عادی"),
    ("هشتگ‌ها", "#CoalTarEpoxy #پوشش_ضدخوردگی #مخازن_دفنی #Painting_Coating"),
]
make_table(meta, [Mm(35), Mm(139)])

# Resources
heading("منابع", border=True)
para("پیوست‌ها (آماده برای بارگذاری، نه بارگذاری‌شده):", bold=True)
para("۱. پیشنهاد تغییر سیستم پوشش مخازن دفنی دیلوج و استوریج کولینگ کمکی.pdf")
para("۲. پیشنهاد تغییر سیستم پوشش مخازن دفنی دیلوج و استوریج کولینگ کمکی.docx")

# QA
heading("وضعیت QA", border=True)
para("نیازمند بازبینی — بدون مسئله حیاتی؛ نیاز به تأیید شجره دانش، پروژه، ادعاهای فنی و کمیته.")

heading("بازبینی اپراتور الزامی است", border=True)
para("بله — موارد زیر نیازمند بازبینی/تأیید است: شجره دانش، پروژه، اعتبار ادعاهای فنی محصول، کمیته.")

heading("موارد حل‌نشده", border=True)
for item in [
    "انتخاب و تأیید نهایی شجره دانش (پیشنهادی: Painting and Coating).",
    "تأیید نام رسمی پروژه (نیروگاه فردوسی/TOUS، قرارداد 3600/19).",
    "تأیید منبع ادعاهای فنی (تأییدیه NACE/API/ISO و پژوهشگاه پلیمر و پتروشیمی ایران).",
    "تأیید کمیته پیشنهادات و بذر پیشنهاد.",
    "[اصلاح احتمالی] قالب شماره مستند مرجع: «MD1-TU-00-GN-W-00-PW2-003» در برابر «MD1/TU-00-GN-W-00-PW2-003» — تأیید اپراتور الزامی است.",
]:
    para("• " + item)

# Checklist table
heading("چک‌لیست نهایی اپراتور", border=True)
checks = ["نوع دانش تأیید شد (پیشنهاد)", "شجره دانش تأیید شد", "پروژه تأیید شد",
          "محدوده سازمانی تأیید شد", "سطح دسترسی تأیید شد", "همکاران تأیید شدند",
          "محتوا بازبینی شد", "پیوست‌ها بازبینی شدند", "هشتگ‌ها بازبینی شدند",
          "مسائل QA حل شدند", "پیش‌نویس نهایی برای ثبت در DANA تأیید شد"]
rows = [("مورد", "وضعیت")] + [(c, "☐") for c in checks]
make_table(rows, [Mm(120), Mm(54)])

# Footer
p = para("تولید شده توسط organizational-knowledge-skill — پیش‌نویس برای بازبینی و تأیید انسانی؛ ثبت نهایی در DANA بر عهده اپراتور است.",
         font=FONT_BODY, size=9, color=RGBColor(0x55, 0x55, 0x55))
p.paragraph_format.space_before = Pt(10)

doc.save(out)
print("OK ->", out)